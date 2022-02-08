from docx import Document
from docx.shared import Inches,Pt,RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
from PIL import Image
import numpy as np
import sys
sys.path.insert(1, 'C:\\Users\HP\python\python project')
from writter import *

wrd=writer()
size=0.8

file=Document("python project/table/table.docx")
img=Image.open("python project/handwritten/images/real.jpg")
a=np.array(img)

table = file.tables[1]
data = []
col=[]
x_dimmension=4500
y_dimmension=5000


for column in table.columns:
    col.append(column.width)
sum_col=sum(col)

div=(x_dimmension-100)/sum_col
arr = np.zeros([y_dimmension,x_dimmension, 3], dtype=np.uint8)
arr[:,:]=[240,240,240]
s=50
col_space=[]
col_add=[80]
for i in col:
    y=int(s+i*div)
    col_space.append(y-s)
    col_add.append(y+30)
    s=y

Dict={}
x=80
col_add[0]=x

l_cell=[]
m_cell=[]
for row in table.rows:
   l=list(cell for cell in row.cells)
   l_cell.append(l)

for row in table.columns:
   l=list(cell for cell in row.cells)
   m_cell.append(l)

#print('m_cell : ',m_cell)

mn=[]
xn=0
for r in range(len(l_cell)-1,-1,-1):
    ln=[]
    for c in range(len(l_cell[0])-1,-1,-1):

        if l_cell[r][c]==l_cell[r][c-1]:
            l_cell[r].pop(c)
            xn+=1
        else:
            ln.insert(0,xn)
            xn=0

        if m_cell[c][r]==m_cell[c][r-1]:
            l_cell[r][c]=False

    ln.insert(0,0)
    mn.insert(0,ln)

for ni,i in enumerate(mn):
    sum=0
    for nj,j in enumerate(i):
        sum+=j
        mn[ni][nj]=sum


print(mn)    
max=[0]*(ni+2)
for m,i1 in enumerate(l_cell): 
    cadd=0
    y=max[m]+200
    y0=y
    x=80
    max[m+1]=y
    for n,cell in enumerate(i1) :
        try:
            if l_cell[m+1][n]==False:
                cadd+=1
        except:
            pass
        if cell==False:
            continue
        y=y0
        for new,para in enumerate(cell.paragraphs) :
            x=col_add[n+mn[m][n]]
            if new:
                y+=170
                if y>max[m+1+cadd]:
                    max[m+1+cadd]=y
            word=para.text.split(' ')
            k=0
            for run in para.runs :
                color=False
                if run.font.color.rgb !=None:
                    color=[]
                    for i in run.font.color.rgb:
                        color.append(i)
                for c in run.text:
                    if c==' ':
                        x+=int(100*size)
                        k=k+1
                        if (col_add[n+1+mn[m][n+1]]-x-((wrd.length(word[k],size)))) < 0:
                                x=col_add[n+mn[m][n]]
                                y+=int(200*size)
                                if y>max[m+1+cadd]:
                                    max[m+1+cadd]=y
                        continue
                    try:
                        pass
                       # x=wrd.write(c,arr,x,y,color,2,size)
                    except:
                        pass
                    
for m,i in enumerate(max) :
    for n in range(len(col_add)-1):
        try:
            if l_cell[m][n]==False:
                continue
        except:
            pass
        arr[i+50:i+60,col_add[n]-30:col_add[n+1]-30]=[83,83,83] 

for n,j in enumerate(l_cell):
    for m in range(len(l_cell[n])+1):
        try:
            arr[max[n]+50:max[n+1]+50,col_add[m+mn[n][m]]-30:n+col_add[m+mn[n][m]]-20]=[83,83,83]       
        except:
            pass

arr=arr[:max[-1]+110,:]
pageimg = Image.fromarray(arr)
pageimg.save("python project/table/table.jpg")
del arr
pageimg.close()
img.close()
wrd.img.close()